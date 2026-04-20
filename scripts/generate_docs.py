import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Skipping {md_path}: File not found")
        return

    print(f"Converting {md_path} to {docx_path}...")
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    for line in lines:
        stripped = line.strip()
        
        # Code Blocks
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
        
        if in_code_block:
            p = doc.add_paragraph(line, style='No Spacing')
            p.paragraph_format.left_indent = Pt(20)
            p.runs[0].font.name = 'Courier New'
            continue

        # Headers
        if stripped.startswith('# '):
            h = doc.add_heading(stripped[2:], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=3)
        
        # Tables (Very basic detection)
        elif stripped.startswith('|'):
            # Convert table row to text
            cells = [c.strip() for c in stripped.split('|') if c.strip()]
            if cells:
                doc.add_paragraph(' | '.join(cells), style='No Spacing')
        
        # Lists
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped.startswith('1. ') or stripped.startswith('2. '):
            # Check if it's a number followed by dot
            import re
            if re.match(r'^\d+\.', stripped):
                doc.add_paragraph(stripped[stripped.find(' ')+1:], style='List Number')
            else:
                doc.add_paragraph(stripped)
        
        # Separator
        elif stripped == '---' or stripped == '***':
            doc.add_paragraph('_' * 40)
        
        # Normal Text
        elif stripped:
            cleaned = stripped.replace('**', '').replace('__', '')
            doc.add_paragraph(cleaned)
        else:
            doc.add_paragraph()

    doc.save(docx_path)
    print(f"Successfully saved: {docx_path}")

def main():
    docs_dir = 'docs'
    output_files = [
        ('PROJECT_ARCHITECTURE.md', 'Project_Architecture.docx'),
        ('DATABASE_DOCUMENTATION.md', 'Database_Documentation.docx'),
        ('API_DOCUMENTATION.md', 'API_Endpoints_and_URLs.docx')
    ]

    for md_file, docx_file in output_files:
        md_path = os.path.join(docs_dir, md_file)
        docx_path = os.path.join(docs_dir, docx_file)
        markdown_to_docx(md_path, docx_path)

if __name__ == "__main__":
    main()
